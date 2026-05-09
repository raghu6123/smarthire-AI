"""
SmartHire AI – Resume Parser
Extracts text and structured information from PDF and DOCX resumes.
"""

import os
import re

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

from utils.constants import SKILL_DICTIONARY, EDUCATION_LEVELS
from utils.helpers import (
    extract_email, extract_phone, extract_name_from_text,
    extract_years_of_experience, clean_text, get_file_extension
)


class ResumeParser:
    """Parses PDF and DOCX resumes to extract structured candidate data."""

    def __init__(self):
        self.supported_formats = ['.pdf', '.docx']

    def parse_resume(self, file_path):
        """
        Parse a resume file and return extracted information.
        Returns a dict with: name, email, phone, skills, education,
        experience, certifications, raw_text
        """
        ext = get_file_extension(file_path)
        if ext == '.pdf':
            raw_text = self._extract_pdf_text(file_path)
        elif ext == '.docx':
            raw_text = self._extract_docx_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        if not raw_text or len(raw_text.strip()) < 10:
            return self._empty_result(file_path)

        return self._extract_info(raw_text, file_path)

    def parse_uploaded_file(self, uploaded_file):
        """Parse a Streamlit UploadedFile object."""
        ext = get_file_extension(uploaded_file.name)
        if ext == '.pdf':
            raw_text = self._extract_pdf_from_bytes(uploaded_file)
        elif ext == '.docx':
            raw_text = self._extract_docx_from_bytes(uploaded_file)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        if not raw_text or len(raw_text.strip()) < 10:
            return self._empty_result(uploaded_file.name)

        return self._extract_info(raw_text, uploaded_file.name)

    def _extract_pdf_text(self, file_path):
        """Extract text from a PDF file on disk."""
        if PyPDF2 is None:
            return ""
        try:
            text = ""
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text
        except Exception as e:
            print(f"Error reading PDF {file_path}: {e}")
            return ""

    def _extract_pdf_from_bytes(self, uploaded_file):
        """Extract text from an uploaded PDF file."""
        if PyPDF2 is None:
            return ""
        try:
            text = ""
            reader = PyPDF2.PdfReader(uploaded_file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text
        except Exception as e:
            print(f"Error reading PDF: {e}")
            return ""

    def _extract_docx_text(self, file_path):
        """Extract text from a DOCX file on disk."""
        if Document is None:
            return ""
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except Exception as e:
            print(f"Error reading DOCX {file_path}: {e}")
            return ""

    def _extract_docx_from_bytes(self, uploaded_file):
        """Extract text from an uploaded DOCX file."""
        if Document is None:
            return ""
        try:
            doc = Document(uploaded_file)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            return ""

    def _extract_info(self, raw_text, source):
        """Extract structured information from raw resume text."""
        cleaned = clean_text(raw_text)
        text_lower = cleaned.lower()

        return {
            "name": extract_name_from_text(raw_text),
            "email": extract_email(raw_text),
            "phone": extract_phone(raw_text),
            "skills": self._extract_skills(text_lower),
            "education": self._extract_education(text_lower),
            "experience": extract_years_of_experience(text_lower),
            "certifications": self._extract_certifications(text_lower),
            "raw_text": cleaned,
            "source": source,
        }

    def _extract_skills(self, text_lower):
        """Extract skills by matching against the skill dictionary."""
        found_skills = []
        for skill in SKILL_DICTIONARY:
            # Use word boundary matching for short skills
            if len(skill) <= 3:
                pattern = r'\b' + re.escape(skill) + r'\b'
                if re.search(pattern, text_lower):
                    found_skills.append(skill)
            else:
                if skill in text_lower:
                    found_skills.append(skill)
        return list(set(found_skills))

    def _extract_education(self, text_lower):
        """Extract highest education level found."""
        found_levels = []
        for level, score in EDUCATION_LEVELS.items():
            if level in text_lower:
                found_levels.append((level, score))
        if found_levels:
            found_levels.sort(key=lambda x: x[1], reverse=True)
            return found_levels[0][0].title()
        return "Not Specified"

    def _extract_certifications(self, text_lower):
        """Extract certifications from text."""
        cert_keywords = [
            "aws certified", "azure certified", "google certified",
            "pmp", "scrum master", "cissp", "ceh", "comptia",
            "cisco certified", "oracle certified", "itil",
            "six sigma", "tensorflow developer", "data engineer",
        ]
        found = [cert for cert in cert_keywords if cert in text_lower]
        return found

    def _empty_result(self, source):
        """Return empty result when parsing fails."""
        return {
            "name": "Unknown",
            "email": "",
            "phone": "",
            "skills": [],
            "education": "Not Specified",
            "experience": 0,
            "certifications": [],
            "raw_text": "",
            "source": source,
        }
